# For the OP: change all EDIT tags back to what you code has set

import os
import re
import random

from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
import openpyxl

from matplotlib import pyplot as plt
from matplotlib import patches as mpatches
from io import BytesIO

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_BREAK

List = []
# EDIT: I changed the path to my own path
wb = openpyxl.load_workbook('schools/schools.xlsx')
sheet = wb['sheet1']
tuple(sheet['A1':'A100'])
for rowOfCellObjects in sheet['A1':'A100']:
    for cellObj in rowOfCellObjects:
        List.append(cellObj.value)

School_list_result = []
State = "NY"

# EDIT: on macos I have env variables set therefore I don't need to pass in an executable param
driver = webdriver.Chrome()


def check_xpath(xpath):
    try:
        element = driver.find_element_by_xpath(xpath)
        School_list_result.append(element.text)
    except NoSuchElementException:
        School_list_result.append("No data.")


def check_text(partial_link_text):
    try:
        element_text = driver.find_element_by_partial_link_text(partial_link_text)
        School_list_result.append(element_text.get_attribute("href"))
    except NoSuchElementException:
        School_list_result.append("No data.")


def check_click(clicker):
    try:
        element_click = driver.find_element_by_partial_link_text(clicker)
        element_click.click()
    except NoSuchElementException:
        print("No click.")


def get_url(url, _xpath, send_keys):
    driver.get(url)
    try:
        _element = driver.find_element_by_xpath(_xpath)
        _element.clear()
        driver.implicitly_wait(10)
        _element.send_keys(schools, send_keys)
        _element.send_keys(u'\ue007')
        driver.implicitly_wait(10)
    except NoSuchElementException:
        print("No data.")


# EDIT: adding code -------------------- start
def extract_grades(arr):
    res = {}
    flag = True
    counter = -1
    ke = []
    for d in arr:
        if len(d) == 0:
            continue
        else:
            if d == 'Ungraded\n':
                res['Ungraded Student'] = ''
                ke.append('Ungraded Student')
            elif d == 'Students':
                flag = False
            elif flag:
                res[d] = ''
                ke.append(d)
            else:
                counter += 1
                val = int(d)
                if val == 0:
                    del res[ke[counter]]
                else:
                    res[ke[counter]] = round(float(d), 1)
    return [list(res.keys()), list(res.values())]


def extract_gender(arr):
    def give_int(string):
        return int(''.join(string.split(',')))
    return [['male', 'female'], [give_int(arr[-2]), give_int(arr[-1])]]


class Color:
    refer = {0: "0", 1: "1", 2: "2", 3: "3", 4: "4", 5: "5", 6: "6", 7: "7", 8: "8", 9: "9", 10: "A",  11: "B",  12: "C",  13: "D",  14: "E",  15: "F"}

    def __init__(self, min_percent, max_percent):
        if min_percent > max_percent:
            raise IndexError('min_percent first param must be less than max_percent second param')
        self.mi = 0.9 * min_percent
        self.ma = 0.9 * max_percent
        self.exist = []

    @staticmethod
    def generate():
        sa = 0
        arr = []
        for i in range(6):
            ind = random.choice([x for x in range(16)])
            arr.append(Color.refer[ind])
            sa += ind
        return sa, arr

    def new(self):
        sa, arr = Color.generate()
        ensure = 0
        while sa < self.mi or sa > self.ma or arr in self.exist:
            ensure += 1
            if ensure > 50:
                raise OverflowError("Please make the difference between min and max percentages larger")
            sa, arr = Color.generate()
        self.exist.append(arr)
        return f"#{''.join(arr)}"


def piechart(race):
    plt.clf()
    plt.style.use('seaborn-talk')
    fig, ax = plt.subplots()
    labels = []
    percents = []
    for i, ra in enumerate(race):
        if i % 2 == 0:
            labels.append(ra)
        else:
            percents.append(round(float(''.join(re.findall(r'[\d]', ra))), 1))
    plt.title("Enrolment by Race/Ethnicity", fontdict={'fontsize': 20})
    patches, texts, auto = ax.pie(percents, labels=labels, autopct='%1.1f%%', shadow=True)
    centre_circle = plt.Circle((0, 0), 0.70, fc='white')
    plt.gcf().gca().add_artist(centre_circle)
    for i in range(len(auto)):
        auto[i].set_fontsize(16)
        texts[i].set_fontsize(15)
    plt.tight_layout()
    file = BytesIO()
    plt.savefig(file)
    return file


def barchart(grades):
    plt.clf()
    c = Color(20, 70)
    colors = []
    labels = grades[0]
    numeric = grades[1]
    for i in range(len(labels)):
        colors.append(c.new())

    def auto_label(vertical_):
        for j, rectangle in enumerate(vertical_):
            height = rectangle.get_height()
            x_value = rectangle.get_x()
            width = rectangle.get_width()
            plt.text(x_value + width / 2., height, "%d" % height, color=colors[j],
                     ha='center', va='bottom')

    vertical = plt.bar(labels, numeric, edgecolor='k', color=colors)
    auto_label(vertical)
    plt.title('Enrolment by Grades')
    plt.gcf().autofmt_xdate()
    file = BytesIO()
    plt.savefig(file)
    return file


def linechart(genders):
    plt.clf()
    fig, ax = plt.subplots()
    sexes = genders[1]
    ax.broken_barh([(0, sexes[0]), (sexes[0], sum(sexes))], [10, 9], facecolors=('#00BFD5', '#FC4F30'))
    ax.set_ylim(5, 15)
    ax.spines['left'].set_visible(False)
    ax.spines['bottom'].set_visible(False)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.set_yticks([10, 25])
    ax.set_axisbelow(True)
    ax.set_yticklabels(['Q1'])
    ax.grid(axis='x')
    ax.text(sexes[0]*0.4, 14.5, "54%", fontsize=18)
    ax.text(sum(sexes), 14.5, "43%", fontsize=18)
    fig.suptitle('Enrolment by Gender (Students Only)', fontsize=16)
    leg1 = mpatches.Patch(color='#00BFD5', label='Male')
    leg2 = mpatches.Patch(color='#FC4F30', label='Female')
    ax.legend(handles=[leg1, leg2], ncol=3, loc='lower right')
    file = BytesIO()
    plt.savefig(file)
    return file


class School:
    @staticmethod
    def dict_last_key(dictionary):
        return list(dictionary.keys())[-1]

    def __init__(self):
        self.headers = {}

    def create_large_header(self, string):
        self.headers[string] = {}

    def create_small_header(self, string):
        self.headers[School.dict_last_key(self.headers)][string] = ''

    def add_info(self, string):
        self.headers[list(self.headers.keys())[-1]][School.dict_last_key(self.headers[list(self.headers.keys())[-1]])] = string

    def add_none(self):
        del self.headers[list(self.headers.keys())[-1]][School.dict_last_key(self.headers[list(self.headers.keys())[-1]])]

    def publish(self):
        tem = {}
        for k, v in self.headers.items():
            if len(v) != 0:
                tem[k] = v
        self.headers = {}
        return tem


class Word:
    def __init__(self, path):
        self.result = []
        self.names = []
        self.schools = []
        path = path.strip()
        if path[0] == '"' or path[0] == "'":
            path = path[1:-1]
        self.path = os.path.abspath(path)

    def add_school(self, name):
        self.schools.append(School())
        self.names.append(name)

    def add_title(self, title):
        self.schools[-1].create_large_header(title)

    def add_header(self, header):
        self.schools[-1].create_small_header(header)

    def add_info(self, info):
        if info == 'No data.':
            self.add_none()
            return
        self.schools[-1].add_info(info)

    def add_piechart(self, race_list):
        if race_list == 'No data.':
            self.add_none()
            return
        self.schools[-1].add_info(piechart(race_list))

    def add_barchart(self, grades):
        if grades == 'No data.':
            self.add_none()
            return
        self.schools[-1].add_info(barchart(grades))

    def add_linechart(self, genders):
        if genders == 'No data.':
            self.add_none()
            return
        self.schools[-1].add_info(linechart(genders))

    def add_none(self):
        self.schools[-1].add_none()

    def publish(self):
        for i, school_name in enumerate(self.names):
            school_info = self.schools[i].publish()
            if len(school_info) != 0:
                self.result.append({school_name: school_info})

    def __str__(self):
        self.publish()
        doc = Document(self.path)
        for school in self.result:
            school_name = list(school.keys())[0]
            school_info = school[school_name]

            page_title = doc.add_paragraph()
            page_title.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            page_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = page_title.add_run(school_name)
            title_run.font.color.rgb = RGBColor(54, 95, 145)
            title_run.font.size = Pt(32)
            title_run.add_break(WD_BREAK.TEXT_WRAPPING)

            for header, info in school_info.items():
                para_header = doc.add_paragraph()
                para_header.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                header_run = para_header.add_run(header)
                header_run.font.size = Pt(28)
                header_run.underline = True
                header_run.add_break()

                for sub_header, sub_info in info.items():
                    para_sub_header = doc.add_paragraph()
                    para_sub_header.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    sub_header_run = para_sub_header.add_run(sub_header)
                    sub_header_run.bold = True
                    sub_header_run.font.size = Pt(23)

                    if isinstance(sub_info, str):
                        para_sub_info = doc.add_paragraph()
                        para_sub_info.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                        sub_info_run = para_sub_info.add_run('\t' + sub_info)
                        sub_info_run.italic = True
                        sub_info_run.font.size = Pt(22)
                        sub_info_run.add_break(WD_BREAK.LINE_CLEAR_ALL)
                    else:
                        doc.add_picture(sub_info, width=Inches(5))
                        doc.add_paragraph('\n')
                doc.add_paragraph('\n')
            doc.add_page_break()
        doc.save(self.path)
        return 'Succeed'


# EDIT: edited saved docx blank file absolute path
docx_path = '/Users/Me/PycharmProjects/Medium/schools/schools.docx'
w = Word(docx_path)

# EDIT: adding code -------------------- end


for schools in List[1:2]:
    # -----------------------------------------GREAT SCHOOLS-------------------------------------------
    get_url("https://www.google.com/", '//*[@id="tsf"]/div[2]/div[1]/div[1]/div/div[2]/input',
            " " + State + " greatschools")
    _clicker = driver.find_element_by_xpath('//*[@id="rso"]/div[1]/div/div[1]/a/h3').click()

    check_xpath('//*[@id="hero"]/div/div[1]/h1')  # School Name

    check_xpath('/html/body/div[6]/div[8]/div/div[1]/div/div/div[2]/div[1]/div[2]/span[1]')  # Principal

    check_text('Principal email')  # Principal’s E-mail

    check_xpath('//*[@id="hero"]/div/div[2]/div[2]/div[3]/div[2]')  # Grade Span

    check_xpath('//*[@id="hero"]/div/div[2]/div[1]/div[1]/div[1]/div[1]/a/div/span[2]')  # Address

    check_xpath('/html/body/div[6]/div[8]/div/div[1]/div/div/div[2]/div[2]/span/a')  # Phone

    check_text('Website')  # Website

    check_xpath('//*[@id="hero"]/div/div[2]/div[1]/div[1]/div[1]/div[2]/a')  # Associations/Communities

    check_xpath('//*[@id="hero"]/div/div[2]/div[2]/div[1]/div/a/div[1]/div')  # GreatSchools Rating

    check_xpath('//*[@id="Students"]/div/div[2]/div[1]/div[2]')  # Enrollment by Race/Ethnicity

    # -----------------------------------------NCES-------------------------------------------

    driver.implicitly_wait(5)
    get_url("https://nces.ed.gov/search/index.asp?q=&btnG=Search#gsc.tab=0", '//*[@id="qt"]', " " + State)
    check_click('Search for Public Schools - ')
    driver.implicitly_wait(5)

    check_xpath('/html/body/div[1]/div[3]/table/tbody/tr[4]/td/table/tbody/tr[7]/td[1]/font[2]')  # School type

    check_xpath('/html/body/div[1]/div[3]/table/tbody/tr[4]/td/table/tbody/tr[7]/td[3]/font')  # Charter

    check_xpath('/html/body/div[1]/div[3]/table/tbody/tr[12]/td/table/tbody/tr[3]/td/table/tbody/tr[2]/td/table/tbody')
    # Enrolment by Gender

    check_xpath(
        '/html/body/div[1]/div[3]/table/tbody/tr[12]/td/table/tbody/tr[1]/td/table/tbody/tr[2]')  # Enrolment by Grade

    # -----------------------------------------USNEWS-------------------------------------------
    driver.implicitly_wait(5)
    url = "https://www.usnews.com/education/best-high-schools/new-york/rankings"
    driver.get(url)
    check_click(schools)
    driver.implicitly_wait(5)

    check_xpath('//*[@id="app"]/div/div/div/div[1]/div/div/div[2]/div[1]/div[2]/p[3]')  # U.S.News Rankings

    # -----------------------------------------PUBLIC REVIEW-------------------------------------------
    driver.implicitly_wait(5)
    get_url("https://www.google.com/", '//*[@id="tsf"]/div[2]/div[1]/div[1]/div/div[2]/input',
            " " + State + " publicschoolreview")
    # clicker = driver.find_element_by_partial_link_text('(2020)').click()
    driver.implicitly_wait(5)

    check_xpath('//*[@id="quick_stats"]/div/div[2]/ul/li[2]/strong')  # Total # Students

    check_xpath('//*[@id="total_teachers_data_row"]/td[2]')  # Full-Time Teachers

    check_xpath('//*[@id="quick_stats"]/div/div[2]/ul/li[3]/strong')  # Student/Teacher Ratio

    # -----------------------------------------PRINT INFOFMATION-------------------------------------------

    print("         ---------------------------------------------------------------" + "\n",
          "                              \033[1m", schools, "\033[0m" + "\n",
          "         ---------------------------------------------------------------" + "\n",
          "                              \033[1mGeneral Information\033[0m        " + "\n",
          "\033[1mSchool Name:\n\033[0m", School_list_result[0] + "\n",
          "\033[1mPrincipal:\n\033[0m", School_list_result[1] + "\n",
          "\033[1mPrincipal’s E-mail:\n\033[0m", School_list_result[2] + "\n",
          "\033[1mType:\n\033[0m", School_list_result[10] + "\n",
          "\033[1mGrade Span:\n\033[0m", School_list_result[3] + "\n",
          "\033[1mAddress:\n\033[0m", School_list_result[4] + "\n",
          "\033[1mPhone:\n\033[0m", School_list_result[5] + "\n",
          "\033[1mWebsite:\n\033[0m", School_list_result[6] + "\n",
          "\033[1mAssociations/Communities:\n\033[0m", School_list_result[7] + "\n",
          "\033[1mGreatSchools Summary Rating:\n\033[0m", School_list_result[8] + "\n",
          "\033[1mU.S.News Rankings:\n\033[0m", School_list_result[14] + "\n",
          "                              \033[1mSchool Details\033[0m" + "\n",
          "\033[1mTotal # Students:\n\033[0m", School_list_result[15] + "\n",
          "\033[1mFull-Time Teachers:\n\033[0m", School_list_result[16] + "\n",
          "\033[1mStudent/Teacher Ratio:\n\033[0m", School_list_result[17] + "\n",
          "\033[1mCharter:\n\033[0m", School_list_result[11] + "\n",
          "\033[1mMagnet: \n\033[0m", "No""\n",
          "                              \033[1mEnrolment Data\033[0m" + "\n",
          "\033[1mEnrolment by Race/Ethnicity: \n\033[0m", School_list_result[9] + "\n",
          "\033[1mEnrolment by Gender: \n\033[0m", School_list_result[12] + "\n",
          "\033[1mEnrolment by Grade: \n\033[0m", School_list_result[13] + "\n",
          ()
          )
    print()

    # EDIT: adding code -------------------- start
    w.add_school(School_list_result[0])

    w.add_title("General Information")
    w.add_header("School Name")
    w.add_info(School_list_result[0])
    w.add_header('Principal')
    w.add_info(School_list_result[1])
    w.add_header('Principal\'s Email')
    w.add_info(School_list_result[2])
    w.add_header('Type')
    w.add_info(School_list_result[10])
    w.add_header('Grade Span')
    w.add_info(School_list_result[3])
    w.add_header('Address')
    w.add_info(School_list_result[4])
    w.add_header('Phone')
    w.add_info(School_list_result[5])
    w.add_header('Website')
    w.add_info(School_list_result[6])
    w.add_header('Associations/Communities')
    w.add_info(School_list_result[7])
    w.add_header('GreatSchools Summary Rating')
    w.add_info(School_list_result[8])
    w.add_header('U.S.News Rankings')
    w.add_info(School_list_result[14])

    w.add_title('School Details')
    w.add_header('Total Students')
    w.add_info(School_list_result[15])
    w.add_header('Full-Time Teachers')
    w.add_info(School_list_result[16])
    w.add_header('Students/Teachers Ratio')
    w.add_info(School_list_result[17])
    w.add_header('Charter')
    w.add_info(School_list_result[11])
    w.add_header('Magnet')
    w.add_info("No")

    w.add_title('Enrolment Data')
    w.add_header('Enrolment by Race/Ethnicity')
    # w.add_info(' '.join(School_list_result[9].split('\n')))
    w.add_piechart(School_list_result[9].split('\n'))
    w.add_header('Enrolment by Gender')
    # w.add_info(School_list_result[12])
    w.add_linechart(extract_gender(School_list_result[12].split(' ')))
    w.add_header('Enrolment by Grade')
    # w.add_info(School_list_result[13])
    w.add_barchart(extract_grades(School_list_result[13].split(' ')))
    School_list_result.clear()
    # EDIT: adding code -------------------- end

print(w)
